#!/usr/bin/env python3

from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Add hyperlink styling
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return hyperlink

# Create document
doc = Document()

# Title
title = doc.add_heading('From Automotive Chaos to AI-Powered Clarity', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('How Infinitra Solved the Universal "First Mile" Problem with Strands Agents')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].italic = True

doc.add_paragraph()

# Section 1
doc.add_heading('The Universal "First Mile" Problem - And What We Discovered', level=1)

doc.add_paragraph('The Pattern We See Everywhere:', style='Intense Quote')
p = doc.add_paragraph()
p.add_run('• Healthcare: ').bold = True
p.add_run('Patient → Receptionist → Nurse → Doctor (context lost at each handoff)')
p = doc.add_paragraph()
p.add_run('• Finance: ').bold = True
p.add_run('Client → Call Center → Advisor → Specialist (starting over each time)')
p = doc.add_paragraph()
p.add_run('• Tech Support: ').bold = True
p.add_run('Customer → Level 1 → Level 2 → Engineer (repeating the same story)')

doc.add_paragraph('In Automotive Repair, We Observed the Real Breakdown:', style='Intense Quote')

doc.add_heading('Communication Chaos During Handoffs:', level=2)
p = doc.add_paragraph()
p.add_run('• Customer describes problem → Service advisor takes notes → Mechanic gets incomplete picture')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('"Most communication breakdowns occur during staging"').bold = True
p.add_run(' - parts status and work coordination fail')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('No institutional memory').bold = True
p.add_run(': Every visit starts from zero, even for repeat customers')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Paper-based chaos').bold = True
p.add_run(': Critical details get lost in handwriting and filing cabinets')

doc.add_heading('The Bottlenecks:', level=2)
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Slow approval processes').bold = True
p.add_run(': Cars sit idle while customers delay repair decisions')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Trust gaps').bold = True
p.add_run(': Customers skeptical without visual evidence of problems')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Reactive maintenance').bold = True
p.add_run(': No proactive scheduling, just crisis management')

highlight = doc.add_paragraph()
highlight.add_run('We Recently Explored This Challenge Using Strands Agents.').bold = True
highlight.add_run(' We experimented with an AI-powered approach that could capture complete context once, preserve it across all handoffs, and build institutional memory that gets smarter with every interaction.')

# Section 2
doc.add_heading('From Chaos to Clarity: Building Institutional Memory', level=1)

p = doc.add_paragraph()
p.add_run('The Core Insight: ').bold = True
p.add_run('Every industry suffers from the same fundamental problem - ')
p.add_run('information decay across human handoffs').bold = True
p.add_run('. What if we could eliminate that decay entirely?')

p = doc.add_paragraph()
p.add_run('Our Automotive Experiment: ').bold = True
p.add_run('Using Strands Agents, we built a system that doesn\'t just store information - it ')
p.add_run('learns from every interaction').bold = True
p.add_run('. Instead of starting fresh each time, the system remembers:')

p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Complete vehicle history').bold = True
p.add_run(': Every past issue, repair, and pattern')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Customer communication style').bold = True
p.add_run(': How they describe problems and preferences')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Diagnostic patterns').bold = True
p.add_run(': What symptoms typically indicate which issues')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Repair outcomes').bold = True
p.add_run(': What worked, what didn\'t, and why')

doc.add_heading('The Breakthrough Approach:', level=2)
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Probing methodology').bold = True
p.add_run(': Ask the right questions to build complete context')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Atomic tools architecture').bold = True
p.add_run(': Each tool does one thing perfectly, AI orchestrates intelligently')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Multi-model consensus').bold = True
p.add_run(': Three AI models validate each other for 95% confidence')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Institutional memory').bold = True
p.add_run(': Every interaction makes the system smarter for everyone')

doc.add_heading('Early Results:', level=2)
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Customers get tentative quotes before visiting').bold = True
p.add_run(' - no wasted trips to the service center')
p = doc.add_paragraph()
p.add_run('• Service advisors instantly access complete vehicle context')
p = doc.add_paragraph()
p.add_run('• Mechanics receive probable diagnoses before cars arrive')

quote = doc.add_paragraph('"We didn\'t waste time figuring it out. We just got to work."')
quote.style = 'Intense Quote'

p = doc.add_paragraph()
p.add_run('The bigger question').bold = True
p.add_run(': If this works for automotive handoffs, what other industries could benefit from institutional memory?')

# Section 3
doc.add_heading('The Architecture Behind the Breakthrough', level=1)

p = doc.add_paragraph()
p.add_run('Why Traditional Approaches Fail: ').bold = True
p.add_run('Most AI implementations try to replace human expertise. We took a different approach - ')
p.add_run('amplify human intelligence').bold = True
p.add_run(' with institutional memory that never forgets.')

doc.add_heading('Strands Agents: The Game-Changing Foundation', level=2)

p = doc.add_paragraph()
p.add_run('The Breakthrough Technology: ').bold = True
p.add_run('Strands Agents revolutionizes how AI systems work by enabling ')
p.add_run('true intelligent tool orchestration').bold = True
p.add_run('. Instead of monolithic AI trying to do everything, Strands agents intelligently decide which tools to use when, based on conversation context.')

p = doc.add_paragraph()
p.add_run('How Strands Changes Everything:').bold = True
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Dynamic tool selection').bold = True
p.add_run(': Agent chooses the right tool for each situation')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Context preservation').bold = True
p.add_run(': Complete conversation history maintained across all tool interactions')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Intelligent orchestration').bold = True
p.add_run(': No rigid workflows - the agent adapts in real-time')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Seamless integration').bold = True
p.add_run(': Tools work together naturally through agent coordination')

doc.add_heading('Atomic Tools Architecture: Powered by Strands', level=2)

p = doc.add_paragraph()
p.add_run('Our 7 Specialized Tools (Built with @tool decorator):').bold = True
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('VIN Decoder').bold = True
p.add_run(': Instantly identifies vehicle specifications')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Symptom Analyzer').bold = True
p.add_run(': Translates customer language into technical terms')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Pattern Matcher').bold = True
p.add_run(': Connects current issues to historical data')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Parts Locator').bold = True
p.add_run(': Finds availability and pricing across suppliers')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Labor Estimator').bold = True
p.add_run(': Calculates time and cost based on vehicle specifics')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Quote Generator').bold = True
p.add_run(': Creates transparent, multi-option estimates')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('History Tracker').bold = True
p.add_run(': Builds comprehensive vehicle profiles over time')

highlight = doc.add_paragraph()
highlight.add_run('The Strands Advantage: ').bold = True
highlight.add_run('The agent ')
highlight.add_run('intelligently orchestrates').bold = True
highlight.add_run(' these tools based on conversation flow, not pre-programmed logic. This creates natural, adaptive interactions that feel human.')

# Continue with remaining sections...
doc.add_heading('Probing Methodology: Strands-Enabled Intelligence', level=2)

p = doc.add_paragraph()
p.add_run('Traditional Approach').bold = True
p.add_run(': "What\'s wrong?" → Generic troubleshooting checklist')
p = doc.add_paragraph()
p.add_run('Strands Approach').bold = True
p.add_run(': Intelligent questioning that builds complete context:')

p = doc.add_paragraph()
p.add_run('1. ').bold = True
p.add_run('Listen first').bold = True
p.add_run(': Strands agent captures customer\'s description')
p = doc.add_paragraph()
p.add_run('2. ').bold = True
p.add_run('Dynamic tool selection').bold = True
p.add_run(': Agent chooses relevant tools based on context')
p = doc.add_paragraph()
p.add_run('3. ').bold = True
p.add_run('Build confidence gradually').bold = True
p.add_run(': Each tool interaction refines diagnostic probability')
p = doc.add_paragraph()
p.add_run('4. ').bold = True
p.add_run('Multi-source validation').bold = True
p.add_run(': Agent orchestrates multiple tools for cross-checking')

p = doc.add_paragraph()
p.add_run('Result').bold = True
p.add_run(': 95% diagnostic confidence through intelligent tool orchestration.')

# Section 4
doc.add_heading('Real-World Results: From Experiment to Production', level=1)

p = doc.add_paragraph()
p.add_run('Performance That Matters: ').bold = True
p.add_run('Our Strands Agents implementation didn\'t just work in theory - it delivered measurable business impact in real-world conditions.')

doc.add_heading('Multi-Model Validation Eliminated Single Points of Failure', level=2)
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Cross-validation').bold = True
p.add_run(' between Amazon Nova Pro and Claude 3.5 Sonnet')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Web search verification').bold = True
p.add_run(' for real-world pricing and repair data')

doc.add_heading('Pattern Recognition Improved with Every Interaction', level=2)
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Institutional memory').bold = True
p.add_run(' that learns from successful diagnoses')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Symptom-to-solution mapping').bold = True
p.add_run(' gets more accurate over time')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Customer communication patterns').bold = True
p.add_run(' optimize future interactions')

doc.add_heading('Customer Experience Transformation', level=2)
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Tentative quotes before shop visits').bold = True
p.add_run(' - no more wasted trips')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Complete vehicle history').bold = True
p.add_run(' available instantly to service advisors')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Transparent pricing').bold = True
p.add_run(' with multiple repair options')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Proactive maintenance').bold = True
p.add_run(' suggestions based on usage patterns')

doc.add_heading('Operational Efficiency', level=2)
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Eliminated communication breakdowns').bold = True
p.add_run(' during service handoffs')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Reduced approval bottlenecks').bold = True
p.add_run(' through clear, evidence-based recommendations')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Streamlined workflows').bold = True
p.add_run(' from customer contact to repair completion')

quote = doc.add_paragraph('"We didn\'t waste time figuring it out. We just got to work."')
quote.style = 'Intense Quote'

p = doc.add_paragraph()
p.add_run('Cost Efficiency').bold = True
p.add_run(': Under $0.000015 per interaction while delivering enterprise-grade intelligence.')

p = doc.add_paragraph()
p.add_run('The Bottom Line').bold = True
p.add_run(': Strands Agents transformed automotive service from reactive problem-solving to proactive, intelligent customer care.')

# Section 5
doc.add_heading('Beyond Automotive: Universal Applications', level=1)

p = doc.add_paragraph()
p.add_run('The Universal Pattern: ').bold = True
p.add_run('Our automotive experiment revealed something profound - ')
p.add_run('every industry suffers from the same "first mile" problem').bold = True
p.add_run('. Anywhere expert knowledge meets customer needs, there\'s an opportunity for institutional memory and intelligent orchestration.')

doc.add_heading('Healthcare Diagnostics', level=2)
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Patient symptoms').bold = True
p.add_run(' → ')
p.add_run('Medical history').bold = True
p.add_run(' → ')
p.add_run('Specialist recommendations').bold = True
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Strands application').bold = True
p.add_run(': Probing methodology that builds complete patient context, remembers medical history, and provides evidence-based treatment suggestions')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Impact').bold = True
p.add_run(': Reduce misdiagnoses, eliminate repeated patient histories, enable proactive care')

doc.add_heading('Financial Advisory', level=2)
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Client goals').bold = True
p.add_run(' → ')
p.add_run('Risk assessment').bold = True
p.add_run(' → ')
p.add_run('Investment strategies').bold = True
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Strands application').bold = True
p.add_run(': Understand client communication style, track financial patterns, provide personalized recommendations')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Impact').bold = True
p.add_run(': Move from generic advice to truly personalized financial planning')

doc.add_heading('Technical Support', level=2)
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Problem description').bold = True
p.add_run(' → ')
p.add_run('System diagnostics').bold = True
p.add_run(' → ')
p.add_run('Solution implementation').bold = True
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Strands application').bold = True
p.add_run(': Build institutional knowledge of common issues, remember customer environments, provide step-by-step guidance')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Impact').bold = True
p.add_run(': Eliminate "have you tried turning it off and on again?" conversations')

doc.add_heading('Legal Consultation', level=2)
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Case details').bold = True
p.add_run(' → ')
p.add_run('Legal research').bold = True
p.add_run(' → ')
p.add_run('Strategy development').bold = True
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Strands application').bold = True
p.add_run(': Pattern recognition across similar cases, comprehensive case history, evidence-based legal strategies')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Impact').bold = True
p.add_run(': Faster case preparation, better client outcomes, reduced research time')

doc.add_heading('The Common Thread', level=2)
p = doc.add_paragraph('Every industry needs:')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Institutional memory').bold = True
p.add_run(' that preserves knowledge across interactions')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Probing methodology').bold = True
p.add_run(' that builds complete context efficiently')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Multi-source validation').bold = True
p.add_run(' that ensures accuracy and builds trust')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Intelligent orchestration').bold = True
p.add_run(' that adapts to each unique situation')

p = doc.add_paragraph()
p.add_run('The Question').bold = True
p.add_run(': Which industry will be next to eliminate their "first mile" problem with Strands Agents?')

# Section 6
doc.add_heading('Infinitra\'s Strands Agents Mastery', level=1)

p = doc.add_paragraph()
p.add_run('Our Strands Journey: ').bold = True
p.add_run('This automotive project wasn\'t just a client delivery - it was our ')
p.add_run('deep dive into Strands Agents technology').bold = True
p.add_run('. As former AWS leaders, we approached it with the same rigor we\'d apply to any mission-critical system.')

doc.add_heading('Key Learnings from Our First Strands Implementation', level=2)

p = doc.add_paragraph()
p.add_run('Atomic Architecture Principles:').bold = True
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Single responsibility tools').bold = True
p.add_run(' work better than monolithic functions')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Let the agent orchestrate').bold = True
p.add_run(' - don\'t hard-code workflows')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Context preservation').bold = True
p.add_run(' is critical for institutional memory')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Tool validation').bold = True
p.add_run(' through multiple sources builds unshakeable confidence')

p = doc.add_paragraph()
p.add_run('Probing Methodology as Universal Framework:').bold = True
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Questions before solutions').bold = True
p.add_run(' works across every domain')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Conversation flow').bold = True
p.add_run(' matters more than rigid scripts')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Customer language').bold = True
p.add_run(' must be preserved and translated, not replaced')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Confidence building').bold = True
p.add_run(' happens gradually through intelligent interaction')

p = doc.add_paragraph()
p.add_run('Production-Ready Patterns:').bold = True
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Feature flags').bold = True
p.add_run(' enable zero-downtime deployment of AI systems')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Multi-model validation').bold = True
p.add_run(' eliminates single points of failure')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Institutional memory').bold = True
p.add_run(' compounds value with every interaction')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Cost optimization').bold = True
p.add_run(' through intelligent prompt engineering (87.2% reduction)')

doc.add_heading('What This Means for Your Industry', level=2)

p = doc.add_paragraph('We\'ve proven these patterns work in the complex, high-stakes world of automotive repair. The same ')
p.add_run('Strands-powered approach').bold = True
p.add_run(' can transform any industry where expert knowledge meets customer needs.')

p = doc.add_paragraph()
p.add_run('Our Advantage').bold = True
p.add_run(': We\'ve already solved the hard problems - tool orchestration, institutional memory, multi-model validation, and production deployment.')

p = doc.add_paragraph()
p.add_run('Your Opportunity').bold = True
p.add_run(': Leverage our proven Strands expertise to eliminate your industry\'s "first mile" problem.')

# Section 7 - Call to Action
doc.add_heading('Ready to Transform Your Industry\'s "First Mile"?', level=1)

p = doc.add_paragraph()
p.add_run('The Opportunity: ').bold = True
p.add_run('Every industry has handoff problems, communication breakdowns, and wasted expertise. We\'ve proven that Strands Agents can solve these challenges with institutional memory and intelligent orchestration.')

p = doc.add_paragraph()
p.add_run('What\'s Possible for You:').bold = True
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Eliminate information decay').bold = True
p.add_run(' across your expert consultation workflows')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Build institutional memory').bold = True
p.add_run(' that gets smarter with every interaction')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Create probing methodologies').bold = True
p.add_run(' tailored to your industry\'s specific needs')
p = doc.add_paragraph()
p.add_run('• ').bold = True
p.add_run('Deploy with confidence').bold = True
p.add_run(' using our proven atomic tools architecture')

p = doc.add_paragraph()
p.add_run('Next Steps: ').bold = True
p.add_run('Whether you\'re in healthcare, finance, legal, or any industry where expertise meets customer needs, we\'d love to explore how Strands Agents could transform your operations.')

p = doc.add_paragraph()
p.add_run('Let\'s start a conversation').bold = True
p.add_run(' about what institutional memory could mean for your business.')

p = doc.add_paragraph()
p.add_run('Contact Infinitra: ').bold = True
contact_p = doc.add_paragraph()
add_hyperlink(contact_p, 'help@theinfinitra.com', 'mailto:help@theinfinitra.com')

# About section
doc.add_paragraph()
doc.add_heading('About Infinitra', level=2)
p = doc.add_paragraph('Infinitra is a technology consultancy founded by former AWS leaders, specializing in cloud-native AI solutions and enterprise transformation. We democratize technology so businesses can innovate, launch, and scale in weeks — not years.')

p = doc.add_paragraph('Learn more at ')
add_hyperlink(p, 'theinfinitra.com', 'https://www.theinfinitra.com')

# Save document
doc.save('/Users/saidachanda/development/dixon-smart-repair/infinitra-strands-agents-blog.docx')
print('DOCX document created successfully!')
