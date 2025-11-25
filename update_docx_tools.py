#!/usr/bin/env python3

from docx import Document

# Load the existing document
doc = Document('/Users/saidachanda/development/dixon-smart-repair/infinitra-strands-agents-blog.docx')

# Find and update the tools section
updated = False
for paragraph in doc.paragraphs:
    if "Our 7 Specialized Tools" in paragraph.text:
        # Update the heading
        paragraph.clear()
        run = paragraph.add_run('Our 5 Core Tools (In Workflow Order):')
        run.bold = True
        updated = True
        break

# Find the tools list and update it
tools_found = False
for i, paragraph in enumerate(doc.paragraphs):
    if not tools_found and paragraph.text and ("VIN Decoder" in paragraph.text or "Symptom Analyzer" in paragraph.text):
        tools_found = True
        
        # Clear the old tools list (next several paragraphs that contain tool descriptions)
        tools_to_clear = []
        for j in range(15):  # Check next 15 paragraphs for tool descriptions
            if i + j < len(doc.paragraphs):
                p_text = doc.paragraphs[i + j].text
                if any(tool in p_text for tool in ["VIN Decoder", "Symptom Analyzer", "Pattern Matcher", "Parts Locator", "Labor Estimator", "Quote Generator", "History Tracker"]):
                    tools_to_clear.append(i + j)
        
        # Clear identified tool paragraphs
        for idx in reversed(tools_to_clear):
            doc.paragraphs[idx].clear()
        
        # Add the new tools list starting from the first cleared paragraph
        if tools_to_clear:
            start_idx = tools_to_clear[0]
            tools_data = [
                ('• ', 'fetch_user_vehicles', ': Check what vehicles the customer has in their account'),
                ('• ', 'extract_vin_with_nova_pro', ': Extract VIN from uploaded images using Amazon Nova Pro'),
                ('• ', 'store_vehicle_record', ': Save new vehicle information to customer\'s account'),
                ('• ', 'search_web', ': Find real-world repair data and pricing using Tavily AI'),
                ('• ', 'calculate_labor_estimates', ': Generate multi-model consensus estimates with validation')
            ]
            
            for j, (bullet, tool_name, description) in enumerate(tools_data):
                if start_idx + j < len(doc.paragraphs):
                    p = doc.paragraphs[start_idx + j]
                    p.add_run(bullet).bold = True
                    p.add_run(tool_name).bold = True
                    p.add_run(description)
        break

# Save the updated document
doc.save('/Users/saidachanda/development/dixon-smart-repair/infinitra-strands-agents-blog-updated.docx')
print('Updated DOCX document with correct 5 tools in workflow order!')
print('Saved as: infinitra-strands-agents-blog-updated.docx')
